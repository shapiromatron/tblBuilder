import abc
import docx
from StringIO import StringIO

from docx import Document


class DOCXReport(object):

    def __init__(self, template_fn, context):
        self.template_fn = template_fn
        self.context = context

    def build_report(self):
        """
        Build DOCX report, create content, return file in StringIO format
        """
        self.doc = Document(self.template_fn)
        self.create_content()

        docx = StringIO()
        self.doc.save(docx)
        docx.seek(0)

        return docx

    @abc.abstractmethod
    def create_content(self):
        """
        Main-method called to generate the content in a Word report
        """
        pass

    def build_table(self, numRows, numCols, widths, cells,
                    numHeaders=1, style=None, firstRowCaption=True):
        """
        Helper function to build a table.

        - numRows: (int)
        - numCols: (int)
        - cells: (list) in the following format:
            [
                {"row":0, "col":0, "text":"value"},
                {"row":1, "col":0, "text":"value", "rowspan": 2},
                {"row":1, "col":0, "text":"value", "colspan": 2},
                {"row":0, "col":0, "runs":[
                    {"text": "standard text"},
                    {"text": "bolded text", "bold": True}
                ]},
            ]

        """

        tbl = self.doc.add_table(rows=numRows, cols=numCols, style=style)

        # set column widths
        tbl.autofit = False
        for i, col in enumerate(tbl.columns):
            col.width = docx.shared.Inches(widths[i])

        # build cells
        for cell in cells:
            cellD = tbl.cell(cell["row"], cell["col"])
            p = cellD.paragraphs[0]

            # merge cells if needed
            rowspan = cell.get("rowspan", None)
            colspan = cell.get("colspan", None)
            if rowspan or colspan:
                rowIdx = cell["row"] + cell.get("rowspan", 1) - 1
                colIdx = cell["col"] + cell.get("colspan", 1) - 1
                cellD.merge(tbl.cell(rowIdx, colIdx))

            # add cell-shading if needed
            if "shade" in cell:
                shade_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
                    docx.oxml.ns.nsdecls('w'), cell["shade"]))
                cellD._tc.get_or_add_tcPr().append(shade_elm)

            # add content
            if "width" in cell:
                cellD.width = docx.shared.Inches(cell["width"])

            if "text" in cell:
                p.text = cell["text"]

            if "runs" in cell:
                for runD in cell["runs"]:
                    run = p.add_run(runD["text"])
                    run.bold = runD.get("bold", False)
                    run.italic = runD.get("italic", False)

        # mark rows as headers to break on pages
        if numHeaders and numHeaders>=1:
            for i in xrange(numHeaders):
                tblHeader = docx.oxml.parse_xml(r'<w:tblHeader {} />'.format(
                    docx.oxml.ns.nsdecls('w')))
                tbl.rows[i]._tr.get_or_add_trPr().append(tblHeader)

        # apply caption-style to the first cell in first-row
        if firstRowCaption:

            cell =  tbl.cell(0, 0)

            cellPr = cell._tc.get_or_add_tcPr()
            cellPr.append(docx.oxml.parse_xml(
                r'<w:tcBorders {} ><w:top w:val="nil"/><w:left w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'.format(
                    docx.oxml.ns.nsdecls('w'))))
            cellPr.append(docx.oxml.parse_xml(
                r'<w:shd {} w:val="clear" w:color="auto" w:fill="auto"/>'.format(
                    docx.oxml.ns.nsdecls('w'))))

            # left-align text using justified (LEFT doesn't work)
            for p in cell.paragraphs:
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW

        return tbl
