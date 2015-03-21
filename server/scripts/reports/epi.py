from docx.shared import Inches
from docx.enum.section import WD_ORIENT

from utils import DOCXReport


def make_landscape(doc):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.page_width  = Inches(11)
    section.page_height  = Inches(8.5)

def build_header_cell(row, col, widths, text, colspan=None, rowspan=None):
    cell = {
        "row": row,
        "col": col,
        "runs": [{ "text": text, "bold": True, "italic": False }]}
    if rowspan:
        cell["rowspan"] = rowspan
    if colspan:
        cell["colspan"] = colspan
        cell["width"] = sum(widths[col:col+colspan])
    else:
        cell["width"] = widths[col]
    return cell

def build_text_cell(row, col, widths, text, rowspan=None, colspan=None):
    cell = {"row": row, "col": col, "text": text}
    if rowspan:
        cell["rowspan"] = rowspan
    if colspan:
        cell["colspan"] = colspan
        cell["width"] = sum(widths[col:col+colspan])
    else:
        cell["width"] = widths[col]
    return cell

def build_run_cell(row, col, widths, runs, rowspan=None, colspan=None):
    cell = {"row": row, "col": col, "runs": runs}
    if rowspan:
        cell["rowspan"] = rowspan
    if colspan:
        cell["colspan"] = colspan
        cell["width"] = sum(widths[col:col+colspan])
    else:
        cell["width"] = widths[col]
    return cell

def run_maker(txt, newline=True, b=False, i=False):
    if newline: txt += u"\n"
    return {"text": txt, "bold": b, "italic": i}

def build_dual_field(row, widths, header, text="", runs=None, colspan=None):
    h = build_header_cell(row, 0, widths, header)
    if runs:
        d = build_run_cell(row, 1, widths, runs, colspan=colspan)
    else:
        d = build_text_cell(row, 1, widths, text, colspan=colspan)
    return [h, d]


class HtmlTblRecreation(DOCXReport):
    """
    Attempt to recreate HTML table in a Word-report.
    """

    def getCol2(self, d):
        # recreation of table-helper
        runs = []
        if d["isCaseControl"]:
            runs.append(run_maker("Cases: ", b=True))
            runs.append(run_maker(u"{} ({}); {}".format(
                            d["populationSizeCase"],
                            d["responseRateCase"],
                            d["sourceCase"])))
            runs.append(run_maker("Controls: ", b=True))
            runs.append(run_maker(u"{} ({}); {}".format(
                        d["populationSizeControl"],
                        d["responseRateControl"],
                        d["sourceControl"])))
        else:
            runs.append(run_maker(u"{}; {}".format(
                        d["populationSize"],
                        d["eligibilityCriteria"])))

        runs.append(run_maker("Exposure assessment method: ", b=True))
        if d["exposureAssessmentType"].lower().find("other")>=0:
            runs.append(run_maker("other", newline=False))
        else:
            runs.append(run_maker(d["exposureAssessmentType"], newline=False))

        if d.get("exposureAssessmentNotes"):
            runs.append(run_maker(u"; " + d.get("exposureAssessmentNotes")))

        if d.get("outcomeDataSource"):
            runs.append(run_maker(d.get("outcomeDataSource"), newline=False))

        return runs

    def build_desc_tbl(self, data):
        rows = 1
        cols = 8
        wds = [1.5, 2.0, 0.8, 1.1, 0.6, 1.0, 0.8, 2.0]

        # write header
        cells = [
            build_header_cell(0, 0, wds, "Reference, location follow-up/enrollment period, study-design"),
            build_header_cell(0, 1, wds, "Population size, description, exposure assessment method"),
            build_header_cell(0, 2, wds, "Organ site (ICD code)"),
            build_header_cell(0, 3, wds, "Exposure category or level"),
            build_header_cell(0, 4, wds, "Exposed cases/ deaths"),
            build_header_cell(0, 5, wds, "Risk estimate (95% CI)"),
            build_header_cell(0, 6, wds, "Covariates controlled"),
            build_header_cell(0, 7, wds, "Comments"),
        ]

        # write additional rows
        for d in data["descriptions"]:

            for res in d["results"]:
                res["_rowspan"] = max(len(res["riskEstimates"]), 1)
                if res["hasTrendTest"]:
                    res["_rowspan"] += 1

            st_rowspan = sum(res["_rowspan"] for res in d["results"])

            # Column A
            runs = [
                run_maker(d["reference"]["name"]),
                run_maker(d["location"]),
                run_maker(d["enrollmentDates"]),
                run_maker(d["studyDesign"]),
            ]
            cells.append(build_run_cell(rows, 0, wds, runs, rowspan=st_rowspan))

            # Column B
            runs = self.getCol2(d)
            cells.append(build_run_cell(rows, 1, wds, runs, rowspan=st_rowspan))

            # Columns C, D, E, F, G
            irows = rows
            for res in d["results"]:
                cells.append(build_text_cell(irows, 2, wds, res["organSite"], rowspan=res["_rowspan"]))
                for i, est in enumerate(res["riskEstimates"]):
                    cells.append(build_text_cell(irows+i, 3, wds, est["exposureCategory"]))
                    cells.append(build_text_cell(irows+i, 4, wds, unicode(est["numberExposed"])))
                    cells.append(build_text_cell(irows+i, 5, wds, unicode(est["riskFormatted"])))

                cells.append(build_text_cell(irows, 6, wds, res["covariatesList"], rowspan=res["_rowspan"]))

                if res["hasTrendTest"]:
                    txt = u"Trend-test p-value: {}".format(res["trendTest"])
                    cells.append(build_text_cell(irows+i+1, 3, wds, txt, colspan=3))

                irows += res["_rowspan"]

            # Column H
            runs = [
                run_maker(d.get("notes")),
                run_maker("Strengths: ", b=True, newline=False),
                run_maker(d.get("strengths")),
                run_maker("Limitations: ", b=True, newline=False),
                run_maker(d.get("limitations")),
            ]
            cells.append(build_run_cell(rows, 7, wds, runs, rowspan=st_rowspan))

            rows += st_rowspan

        self.build_table(rows, cols, wds, cells, numHeaders=1, firstRowCaption=False, style="ntpTbl")

    def create_content(self):
        doc = self.doc
        d = self.context

        make_landscape(doc)

        # title
        txt = "{} {}: {}".format(
            d["table"]["volumeNumber"],
            d["table"]["monographAgent"],
            d["table"]["name"]
        )
        doc.add_heading(txt, 0)
        self.build_desc_tbl(d)


class NtpEpiDescriptive(DOCXReport):

    def build_desc_tbl(self, d):
        rows = 2
        cols = 5
        wds = [1.5, 0.7, 1.4, 1.4, 5.0]

        # write header
        txt = u"Table X: Study description and methodologies: {}".format(d["reference"]["name"])
        cells = [
            build_header_cell(0, 0, wds, txt, colspan=5),
            build_header_cell(1, 0, wds, "Field"),
            build_header_cell(1, 1, wds, "Description", colspan=4),
            build_header_cell(1, 2, wds, ""),
            build_header_cell(1, 3, wds, ""),
            build_header_cell(1, 4, wds, ""),
        ]

        # write data-rows
        runs = [
            run_maker(d["reference"]["name"], i=True),
            run_maker(d["reference"]["fullCitation"], newline=False)
        ]
        cells.extend(build_dual_field(rows, wds, "Reference", runs=runs, colspan=4))
        rows += 1

        txt = d["studyDesign"]
        cells.extend(build_dual_field(rows, wds, "Study-design type", text=txt, colspan=4))
        rows += 1

        txt = u"{}; {}".format(d["location"], d["enrollmentDates"])
        cells.extend(build_dual_field(rows, wds, "Location and enrollment dates", text=txt, colspan=4))
        rows += 1

        txt = d.get("populationDescription", "")
        cells.extend(build_dual_field(rows, wds, "Population description", text=txt, colspan=4))
        rows += 1

        if d["isCaseControl"]:

            # build specialized table
            txt = "Case-control description and eligibility criteria"
            cells.append(build_header_cell(rows, 0, wds, txt, rowspan=3))

            cells.append(build_text_cell(rows, 1, wds, " "))
            cells.append(build_run_cell(rows, 2, wds, [run_maker("Population size", i=True, newline=False)]))
            cells.append(build_run_cell(rows, 3, wds, [run_maker("Response rates", i=True, newline=False)]))
            cells.append(build_run_cell(rows, 4, wds, [run_maker("Source", i=True, newline=False)]))
            rows +=1

            cells.append(build_text_cell(rows, 1, wds, "Cases"))
            cells.append(build_text_cell(rows, 2, wds, d["populationSizeCase"]))
            cells.append(build_text_cell(rows, 3, wds, d["responseRateCase"]))
            cells.append(build_text_cell(rows, 4, wds, d["sourceCase"]))
            rows +=1

            cells.append(build_text_cell(rows, 1, wds, "Controls"))
            cells.append(build_text_cell(rows, 2, wds, d["populationSizeControl"]))
            cells.append(build_text_cell(rows, 3, wds, d["responseRateControl"]))
            cells.append(build_text_cell(rows, 4, wds, d["sourceControl"]))
            rows +=1

        else:
            txt = d.get("eligibilityCriteria", "")
            cells.extend(build_dual_field(rows, wds, "Eligibility criteria", text=txt, colspan=4))
            rows += 1

            runs = [
                run_maker("Population size: ", i=True, newline=False),
                run_maker(unicode(d.get("populationSize", ""))),
                run_maker("Loss-to-follow-up: ", i=True, newline=False),
                run_maker(unicode(d.get("lossToFollowUp", ""))),
                run_maker("Referent Group: ", i=True, newline=False),
                run_maker(d.get("referentGroup", ""), newline=False),
            ]
            cells.extend(build_dual_field(rows, wds, "Cohort details", runs=runs, colspan=4))
            rows += 1

            txt = d.get("outcomeDataSource", "")
            cells.extend(build_dual_field(rows, wds, "Outcome data source", text=txt, colspan=4))
            rows += 1

        txt = d["exposureAssessmentType"]
        cells.extend(build_dual_field(rows, wds, "Exposure assessment", text=txt, colspan=4))
        rows += 1

        txt = d.get("exposureAssessmentNotes", "")
        cells.extend(build_dual_field(rows, wds, "Exposure assessment notes", text=txt, colspan=4))
        rows += 1

        txt = d.get("exposureLevel", "")
        cells.extend(build_dual_field(rows, wds, "Exposure-level", text=txt, colspan=4))
        rows += 1

        txt = d.get("coexposuresList", "")
        cells.extend(build_dual_field(rows, wds, "Coexposures", text=txt, colspan=4))
        rows += 1

        covariates = []
        ccTxt = set()
        for res in d.get("results", []):
            covariates.extend(res["covariates"])
            txt = res.get("covariatesControlledText")
            if txt:
                ccTxt.add(txt)
        covariates = u", ".join(sorted(set(covariates)))
        ccTxt = u", ".join(sorted(ccTxt))

        runs = [
            run_maker("Analytical methods: ", i=True, newline=False),
            run_maker(d.get("notes", "")),
            run_maker("Covariates: ", i=True, newline=False),
            run_maker(covariates),
            run_maker("Confounder consideration: ", i=True, newline=False),
            run_maker(ccTxt, newline=False),
        ]
        cells.extend(build_dual_field(rows, wds, "Analysis methods and control for confounding", runs=runs, colspan=4))
        rows += 1

        self.build_table(rows, cols, wds, cells, numHeaders=2, style="ntpTbl")

    def create_content(self):
        doc = self.doc
        d = self.context

        make_landscape(doc)

        # title
        txt = "{} {}: Study descriptions and methodologies".format(d["volumeNumber"], d["monographAgent"])
        doc.add_heading(txt, 0)

        # build table for each organ-site
        for desc in d["descriptions"]:
            self.build_desc_tbl(desc)
            self.doc.add_page_break()


class NtpEpiResults(DOCXReport):

    def build_res_tbl(self, caption, results):
        rows = 2
        cols = 7
        wds = [1.5, 1.5, 0.75, 0.75, 1.0, 1.5, 3.0]

        # write header
        cells = [
            build_header_cell(0, 0, wds, caption, colspan=7),
            build_header_cell(1, 0, wds, "Reference, study-design, location, and year"),
            build_header_cell(1, 1, wds, "Population description & exposure assessment method"),
            build_header_cell(1, 2, wds, "Exposure category or level"),
            build_header_cell(1, 3, wds, "Exposed cases/deaths"),
            build_header_cell(1, 4, wds, "Risk estimate\n(95% CI)"),
            build_header_cell(1, 5, wds, "Co-variates controlled"),
            build_header_cell(1, 6, wds, "Comments, strengths, and weaknesses"),
        ]

        # write additional rows
        for res in results:

            rowspan = len(res["riskEstimates"])
            if res["hasTrendTest"]:
                rowspan += 1

            # Column A
            txt = u"{}\n{}\n{}\n{}".format(
                res["descriptive"]["reference"]["name"],
                res["descriptive"]["studyDesign"],
                res["descriptive"]["location"],
                res["descriptive"]["enrollmentDates"]
            )
            cells.append(build_text_cell(rows, 0, wds, txt, rowspan=rowspan))

            # Column B
            runs = [
                run_maker(res["descriptive"].get("eligibilityCriteria", "")),
                run_maker("Exposure assessment method: ", b=True, newline=False),
                run_maker(res["descriptive"]["exposureAssessmentType"], newline=False)
            ]
            cells.append(build_run_cell(rows, 1, wds, runs, rowspan=rowspan))

            # Columns C,D,E
            for i, est in enumerate(res["riskEstimates"]):
                cells.append(build_text_cell(rows+i, 2, wds, est["exposureCategory"]))
                cells.append(build_text_cell(rows+i, 3, wds, unicode(est["numberExposed"])))
                cells.append(build_text_cell(rows+i, 4, wds, unicode(est["riskFormatted"])))

            if res["hasTrendTest"]:
                txt = u"Trend-test p-value: {}".format(res["trendTest"])
                cells.append(build_text_cell(rows+i+1, 2, wds, txt, colspan=3))

            # Column F
            txt = res["covariatesList"]
            cells.append(build_text_cell(rows, 5, wds, txt, rowspan=rowspan))

            # Column G
            runs = [
                run_maker(res["descriptive"].get("exposureLevel", "")),
                run_maker("Confounding:", b=True),
                run_maker(res.get("covariatesControlledText", "")),
                run_maker("Strengths:", b=True),
                run_maker(res["descriptive"]["strengths"]),
                run_maker("Limitations:", b=True),
                run_maker(res["descriptive"]["limitations"], newline=False)
            ]
            cells.append(build_run_cell(rows, 6, wds, runs, rowspan=rowspan))
            rows += rowspan

        self.build_table(rows, cols, wds, cells, numHeaders=2, style="ntpTbl")

    def create_content(self):
        doc = self.doc
        d = self.context

        make_landscape(doc)

        # title
        txt = "{} {}: Results by organ-site".format(d["volumeNumber"], d["monographAgent"])
        doc.add_heading(txt, 0)

        # build table for each organ-site
        for organSite in d["organSites"]:
            txt = "Table X: Epidemiological exposure to {}: {}".format(d["monographAgent"], organSite["organSite"])
            self.build_res_tbl(txt, organSite["results"])
            self.doc.add_page_break()


class NtpEpiAniResults(DOCXReport):

    def build_res_tbl(self, caption, studies):
        rows = 2
        cols = 7
        wds = [1.5, 1.5, 1.5, 1.0, 1.0, 1.0, 2.5]

        # write header
        cells = [
            build_header_cell(0, 0, wds, caption, colspan=7),
            build_header_cell(1, 0, wds, "Reference & year, animal, study duration"),
            build_header_cell(1, 1, wds, "Substance & purity"),
            build_header_cell(1, 2, wds, "Dosing regimen"),
            build_header_cell(1, 3, wds, "Dose levels"),
            build_header_cell(1, 4, wds, "# animals at sacrifice"),
            build_header_cell(1, 5, wds, "Tumor incidence (n/N) (%)"),
            build_header_cell(1, 6, wds, "Comments, strengths, and limitations"),
        ]

        # write additional rows
        for st in studies:

            # get row-span for each study
            for res in st["results"]:
                res["_rowspan"] = len(res["riskEstimates"])+1
                if res["hasTrendTest"]:
                    res["_rowspan"] += 1

            st_rowspan = sum(res["_rowspan"] for res in st["results"])

            # Column A
            runs = [
                run_maker(st["reference"]["name"], b=True),
                run_maker(st.get("eligibilityCriteria", "")),
                run_maker(st.get("enrollmentDates", "")),
            ]
            cells.append(build_run_cell(rows, 0, wds, runs, rowspan=st_rowspan))

            # Column B
            runs = [
                run_maker(st["location"]),
                run_maker(st["coexposuresList"]),
            ]
            cells.append(build_run_cell(rows, 1, wds, runs, rowspan=st_rowspan))

            # Column C
            runs = [
                run_maker(st.get("exposureAssessmentNotes", ""))
            ]
            cells.append(build_run_cell(rows, 2, wds, runs, rowspan=st_rowspan))

            # Columns D, E, F
            irows = rows
            covariates = []
            for res in st["results"]:
                runs = [
                    run_maker(res.get("covariatesControlledText", ""), b=True, newline=False),
                ]
                cells.append(build_run_cell(irows, 3, wds, runs, colspan=3))

                for i, est in enumerate(res["riskEstimates"]):
                    cells.append(build_text_cell(irows+i+1, 3, wds, est["exposureCategory"]))
                    cells.append(build_text_cell(irows+i+1, 4, wds, unicode(est["numberExposed"])))
                    txt = u"{}/{} ({}%)".format(est["riskLow"], est["riskHigh"], est["riskMid"])
                    cells.append(build_text_cell(irows+i+1, 5, wds, txt))

                if res["hasTrendTest"]:
                    txt = u"Trend-test p-value: {}".format(res["trendTest"])
                    cells.append(build_text_cell(irows+i+2, 3, wds, txt, colspan=3))

                covariates.extend(res["covariates"])
                irows += res["_rowspan"]

            # Column G
            covariates = u", ".join(set(covariates))
            runs = [
                run_maker(covariates),
                run_maker("Strengths: ", b=True, newline=False),
                run_maker(st["strengths"]),
                run_maker("Limitations: ", b=True, newline=False),
                run_maker(st["limitations"]),
                run_maker("Other comments: ", b=True, newline=False),
                run_maker(st["notes"], newline=False),
            ]
            cells.append(build_run_cell(rows, 6, wds, runs, rowspan=st_rowspan))

            rows += st_rowspan

        self.build_table(rows, cols, wds, cells, numHeaders=2, style="ntpTbl")

    def create_content(self):
        doc = self.doc
        d = self.context

        make_landscape(doc)

        # title
        txt = "{} {}: Results by organ-site".format(d["volumeNumber"], d["monographAgent"])
        doc.add_heading(txt, 0)

        # build table for each organ-site
        for organSite in d["organSites"]:
            txt = "Table X: Animal-bioassay exposure to {}: {}".format(d["monographAgent"], organSite["organSite"])
            self.build_res_tbl(txt, organSite["studies"])
            self.doc.add_page_break()
